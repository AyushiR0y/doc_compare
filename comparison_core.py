from __future__ import annotations

import base64
import difflib
import logging
import re
import string
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF
from docx import Document
from PIL import Image, ImageDraw

# Enable debug logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)


ALLOWED_EXTENSIONS = {"pdf", "docx"}
_PUNCT_TRANSLATOR = str.maketrans("", "", string.punctuation)
_LOW_SIGNAL_WORDS = {
    "a",
    "an",
    "and",
    "as",
    "at",
    "by",
    "for",
    "from",
    "in",
    "is",
    "it",
    "of",
    "on",
    "or",
    "the",
    "to",
    "with",
}


def _normalize_word(word: str) -> str:
    return word.translate(_PUNCT_TRANSLATOR).lower().strip()


def _collect_word_tokens(word_objects: list[dict[str, Any]]) -> list[str]:
    return [obj["text"] for obj in word_objects if obj.get("type") == "word" and obj.get("text")]


def _is_low_signal_token(token: str) -> bool:
    return len(token) <= 1 or token in _LOW_SIGNAL_WORDS


def _filter_isolated_diff_indices(indices: set[int], support_window: int = 3, min_neighbors: int = 1) -> set[int]:
    if not indices:
        return set()

    sorted_indices = sorted(indices)
    filtered: set[int] = set()

    for pos, idx in enumerate(sorted_indices):
        neighbors = 0

        left = pos - 1
        while left >= 0 and idx - sorted_indices[left] <= support_window:
            neighbors += 1
            left -= 1

        right = pos + 1
        while right < len(sorted_indices) and sorted_indices[right] - idx <= support_window:
            neighbors += 1
            right += 1

        if neighbors >= min_neighbors:
            filtered.add(idx)

    return filtered


def _convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX to PDF using Word COM automation on Windows."""
    logger.info("Starting DOCX to PDF conversion...")
    try:
        import win32com.client
    except ImportError:
        raise RuntimeError("pywin32 is required for DOCX-to-PDF conversion. Install with: pip install pywin32")

    # Create temporary files for input and output
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
        tmp_docx.write(docx_bytes)
        tmp_docx_path = tmp_docx.name
    logger.debug(f"Created temp DOCX file: {tmp_docx_path}")

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
        tmp_pdf_path = tmp_pdf.name
    logger.debug(f"Created temp PDF file: {tmp_pdf_path}")

    try:
        # Get Word COM object
        logger.debug("Dispatching Word.Application...")
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        logger.debug("Word application initialized")

        try:
            # Open the DOCX file
            logger.debug(f"Opening DOCX: {tmp_docx_path}")
            doc = word.Documents.Open(str(Path(tmp_docx_path).resolve()))
            logger.debug("DOCX opened successfully")

            # Export as PDF (wdFormatPDF = 17)
            logger.debug(f"Saving to PDF: {tmp_pdf_path}")
            doc.SaveAs2(str(Path(tmp_pdf_path).resolve()), FileFormat=17)
            doc.Close(0)
            logger.debug("Document saved and closed")

            # Read the PDF bytes
            with open(tmp_pdf_path, "rb") as f:
                pdf_bytes = f.read()
            logger.info(f"Conversion successful. PDF size: {len(pdf_bytes)} bytes")

            return pdf_bytes
        finally:
            logger.debug("Quitting Word application...")
            word.Quit()
    finally:
        # Clean up temporary files
        try:
            Path(tmp_docx_path).unlink()
        except Exception:
            pass
        try:
            Path(tmp_pdf_path).unlink()
        except Exception:
            pass


def _extract_words_from_word(file_bytes: bytes) -> tuple[str, list[dict[str, Any]], Document]:
    logger.debug("[WORD] Starting DOCX word extraction...")
    doc = Document(BytesIO(file_bytes))

    word_objects: list[dict[str, Any]] = []
    text_segments: list[str] = []

    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    def process_paragraph(para: Paragraph) -> None:
        if not para.text.strip():
            return

        is_heading_para = para.style.name.startswith("Heading") if para.style else False
        tokens = para.text.split()
        for token in tokens:
            if not token.strip():
                continue

            is_bold = False
            is_italic = False
            for run in para.runs:
                if token in run.text:
                    if run.bold:
                        is_bold = True
                    if run.italic:
                        is_italic = True
                    break

            word_objects.append(
                {
                    "text": token,
                    "type": "word",
                    "is_bold": is_bold,
                    "is_italic": is_italic,
                    "is_heading": is_heading_para,
                    "in_table": False,
                }
            )
            text_segments.append(token)

        word_objects.append({"type": "newline", "text": "\n"})

    def add_inline_images(run, paragraph_is_table: bool) -> None:
        image_rel_ids = run._element.xpath('.//a:blip/@r:embed')
        for rel_id in image_rel_ids:
            image_part = doc.part.related_parts.get(rel_id)
            if not image_part:
                continue

            word_objects.append(
                {
                    "type": "image",
                    "src": base64.b64encode(image_part.blob).decode("utf-8"),
                    "content_type": image_part.content_type,
                    "in_table": paragraph_is_table,
                }
            )

    def process_table(table: Table) -> None:
        word_objects.append({"type": "table_start", "text": ""})

        for row in table.rows:
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    for token in cell_text.split():
                        if not token.strip():
                            continue

                        word_objects.append(
                            {
                                "text": token,
                                "type": "word",
                                "is_bold": False,
                                "is_italic": False,
                                "is_heading": False,
                                "in_table": True,
                            }
                        )
                        text_segments.append(token)

                for para in cell.paragraphs:
                    for run in para.runs:
                        add_inline_images(run, True)

                if cell_idx < len(row.cells) - 1:
                    word_objects.append({"type": "separator", "text": "|"})

            word_objects.append({"type": "row_end", "text": "\n"})

        word_objects.append({"type": "table_end", "text": ""})
        word_objects.append({"type": "newline", "text": "\n"})

    for element in doc.element.body:
        if isinstance(element, CT_P):
            para = Paragraph(element, doc)
            process_paragraph(para)
            for run in para.runs:
                add_inline_images(run, False)
        elif isinstance(element, CT_Tbl):
            process_table(Table(element, doc))

    logger.debug(f"[WORD] Extracted {len(text_segments)} words from DOCX")
    logger.debug(f"[WORD] First 20 tokens: {text_segments[:20]}")
    return " ".join(text_segments), word_objects, doc


def _extract_words_from_pdf(file_bytes: bytes) -> tuple[str, list[dict[str, Any]], list[dict[str, Any]], fitz.Document]:
    doc = fitz.open(stream=file_bytes, filetype="pdf")

    word_objects: list[dict[str, Any]] = []
    highlight_data: list[dict[str, Any]] = []
    text_segments: list[str] = []

    for page_num in range(len(doc)):
        page = doc[page_num]

        table_rects: list[tuple[float, float, float, float]] = []
        try:
            tables = page.find_tables()
            if tables:
                for table in tables:
                    table_rects.append(table.bbox)
        except Exception:
            table_rects = []

        words = page.get_text("words")
        prev_block = -1

        for word_info in words:
            x0, y0, x1, y1, word_text = word_info[:5]
            block_no = word_info[5]

            if block_no != prev_block:
                if prev_block != -1:
                    word_objects.append({"type": "newline", "text": "\n"})
                prev_block = block_no

            if not word_text.strip():
                continue

            in_table = False
            for table_rect in table_rects:
                if x0 >= table_rect[0] and x1 <= table_rect[2] and y0 >= table_rect[1] and y1 <= table_rect[3]:
                    in_table = True
                    break

            word = word_text.strip()
            word_objects.append(
                {
                    "text": word,
                    "type": "word",
                    "is_bold": False,
                    "is_italic": False,
                    "is_heading": False,
                    "in_table": in_table,
                }
            )
            highlight_data.append({"text": word, "bbox": [x0, y0, x1, y1], "page": page_num, "in_table": in_table})
            text_segments.append(word)

        word_objects.append({"type": "newline", "text": "\n"})

    logger.debug(f"[PDF] Extracted {len(text_segments)} words from PDF")
    logger.debug(f"[PDF] First 20 tokens: {text_segments[:20]}")
    return " ".join(text_segments), word_objects, highlight_data, doc


def _run_diff(words1: list[str], words2: list[str]) -> tuple[set[int], set[int], dict[str, int]]:
    logger.debug(f"[DIFF] Document 1: {len(words1)} words")
    logger.debug(f"[DIFF] Document 2: {len(words2)} words")
    
    # Show first/last 10 words from each for debugging
    logger.debug(f"[DIFF] Doc1 first 10: {words1[:10]}")
    logger.debug(f"[DIFF] Doc1 last 10: {words1[-10:]}")
    logger.debug(f"[DIFF] Doc2 first 10: {words2[:10]}")
    logger.debug(f"[DIFF] Doc2 last 10: {words2[-10:]}")
    
    norm_words1: list[str] = []
    norm_words2: list[str] = []
    norm_to_orig_idx1: list[int] = []
    norm_to_orig_idx2: list[int] = []

    for orig_idx, word in enumerate(words1):
        normalized = _normalize_word(word)
        if normalized:
            norm_words1.append(normalized)
            norm_to_orig_idx1.append(orig_idx)

    for orig_idx, word in enumerate(words2):
        normalized = _normalize_word(word)
        if normalized:
            norm_words2.append(normalized)
            norm_to_orig_idx2.append(orig_idx)

    logger.debug(f"[DIFF] After normalization: Doc1={len(norm_words1)} words, Doc2={len(norm_words2)} words")
    logger.debug(f"[DIFF] Doc1 normalized first 10: {norm_words1[:10]}")
    logger.debug(f"[DIFF] Doc2 normalized first 10: {norm_words2[:10]}")

    matcher = difflib.SequenceMatcher(None, norm_words1, norm_words2, autojunk=False)
    opcodes = matcher.get_opcodes()

    diff_indices1: set[int] = set()
    diff_indices2: set[int] = set()

    # Work in normalized-index space first, then map back to original token indices.
    diff_norm_indices1: set[int] = set()
    diff_norm_indices2: set[int] = set()
    for tag, i1, i2, j1, j2 in opcodes:
        if tag in {"replace", "delete"}:
            for idx in range(i1, i2):
                if not _is_low_signal_token(norm_words1[idx]):
                    diff_norm_indices1.add(idx)
        if tag in {"replace", "insert"}:
            for idx in range(j1, j2):
                if not _is_low_signal_token(norm_words2[idx]):
                    diff_norm_indices2.add(idx)

    diff_norm_indices1 = _filter_isolated_diff_indices(diff_norm_indices1)
    diff_norm_indices2 = _filter_isolated_diff_indices(diff_norm_indices2)

    diff_indices1.update(norm_to_orig_idx1[idx] for idx in diff_norm_indices1)
    diff_indices2.update(norm_to_orig_idx2[idx] for idx in diff_norm_indices2)

    logger.debug(f"[DIFF] Result: Doc1 has {len(diff_indices1)} different words, Doc2 has {len(diff_indices2)} different words")
    
    # Log first 30 different words from each document for debugging
    diff_words1_sample = sorted(diff_indices1)[:30]
    diff_words2_sample = sorted(diff_indices2)[:30]
    logger.debug(f"[DIFF] Doc1 first 30 different word indices: {diff_words1_sample}")
    logger.debug(f"[DIFF] Doc1 first 30 different words: {[words1[i] for i in diff_words1_sample if i < len(words1)]}")
    logger.debug(f"[DIFF] Doc2 first 30 different word indices: {diff_words2_sample}")
    logger.debug(f"[DIFF] Doc2 first 30 different words: {[words2[i] for i in diff_words2_sample if i < len(words2)]}")
    
    total_matching = sum(i2 - i1 for tag, i1, i2, _, _ in opcodes if tag == "equal")
    info = {
        "total_matching": total_matching,
        "total_words1": len(words1),
        "total_words2": len(words2),
        "diff_words1": len(diff_indices1),
        "diff_words2": len(diff_indices2),
    }

    return diff_indices1, diff_indices2, info


def _highlight_pdf_words(doc: fitz.Document, word_data: list[dict[str, Any]], diff_indices: set[int]) -> BytesIO:
    highlighted_doc = fitz.open()
    try:
        page_to_words: dict[int, list[dict[str, Any]]] = {}
        for word_idx in diff_indices:
            if word_idx >= len(word_data):
                continue
            word_info = word_data[word_idx]
            page_to_words.setdefault(word_info["page"], []).append(word_info)

        for page_num in range(len(doc)):
            page = doc[page_num]
            new_page = highlighted_doc.new_page(width=page.rect.width, height=page.rect.height)
            new_page.show_pdf_page(new_page.rect, doc, page_num)

            for word_info in page_to_words.get(page_num, []):

                bbox = word_info["bbox"]
                rect = fitz.Rect(bbox[0], bbox[1], bbox[2], bbox[3])
                try:
                    highlight = new_page.add_highlight_annot(rect)
                    if word_info.get("in_table", False):
                        highlight.set_colors(stroke=[1.0, 0.93, 0.88])
                    else:
                        highlight.set_colors(stroke=fitz.utils.getColor("yellow"))
                    highlight.update()
                except Exception:
                    continue

        output = BytesIO()
        highlighted_doc.save(output)
        output.seek(0)
        return output
    finally:
        highlighted_doc.close()


def _highlight_word_doc(doc: Document, word_objects: list[dict[str, Any]], diff_indices: set[int]) -> BytesIO:
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

    text_idx = 0
    obj_indices_to_highlight: set[int] = set()
    for obj_idx, obj in enumerate(word_objects):
        if obj["type"] == "word":
            if text_idx in diff_indices:
                obj_indices_to_highlight.add(obj_idx)
            text_idx += 1

    current_obj_idx = 0

    def copy_run_style(src: Run, dst: Run) -> None:
        dst.style = src.style
        dst.bold = src.bold
        dst.italic = src.italic
        dst.underline = src.underline
        dst.font.name = src.font.name
        dst.font.size = src.font.size
        dst.font.color.rgb = src.font.color.rgb

    def apply_token_level_highlight(para: Paragraph, color: WD_COLOR_INDEX) -> None:
        nonlocal current_obj_idx

        original_runs = list(para.runs)
        for run in original_runs:
            if not run.text:
                continue

            parts = re.findall(r"\S+|\s+", run.text)
            segmented: list[tuple[str, bool]] = []
            has_highlight = False

            for part in parts:
                if part.isspace():
                    segmented.append((part, False))
                    continue

                while current_obj_idx < len(word_objects) and word_objects[current_obj_idx]["type"] != "word":
                    current_obj_idx += 1

                should_highlight = current_obj_idx < len(word_objects) and current_obj_idx in obj_indices_to_highlight
                segmented.append((part, should_highlight))
                has_highlight = has_highlight or should_highlight
                current_obj_idx += 1

            if not has_highlight:
                continue

            first_text, first_highlight = segmented[0]
            run.text = first_text
            run.font.highlight_color = color if first_highlight else None

            prev_run = run
            for text_part, should_highlight in segmented[1:]:
                new_run = para.add_run(text_part)
                copy_run_style(run, new_run)
                new_run.font.highlight_color = color if should_highlight else None
                prev_run._r.addnext(new_run._r)
                prev_run = new_run

    for element in doc.element.body:
        if isinstance(element, CT_P):
            para = Paragraph(element, doc)
            if not para.text.strip():
                continue
            apply_token_level_highlight(para, WD_COLOR_INDEX.YELLOW)

        elif isinstance(element, CT_Tbl):
            table = Table(element, doc)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        apply_token_level_highlight(para, WD_COLOR_INDEX.BRIGHT_GREEN)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _create_html_preview(word_objects: list[dict[str, Any]], diff_indices: set[int]) -> tuple[str, int]:
    import html

    # Split word-based HTML into page-like blocks for clearer page-wise preview.
    WORDS_PER_PAGE = 350

    pages: list[str] = []
    current_parts: list[str] = []
    paragraph_tokens: list[str] = []
    paragraph_is_heading = False
    text_idx = 0
    in_table = False
    words_in_current_page = 0

    def flush_paragraph() -> None:
        nonlocal words_in_current_page
        nonlocal paragraph_is_heading
        if paragraph_tokens:
            if paragraph_is_heading:
                current_parts.append('<h3 class="word-preview-heading">')
                current_parts.append(" ".join(paragraph_tokens))
                current_parts.append('</h3>')
            else:
                current_parts.append('<p class="word-preview-paragraph">')
                current_parts.append(" ".join(paragraph_tokens))
                current_parts.append('</p>')
            paragraph_tokens.clear()
            paragraph_is_heading = False

    def flush_table_end():
        current_parts.append('</td></tr></tbody></table></div>')

    def start_new_page_if_needed():
        nonlocal words_in_current_page, current_parts
        if words_in_current_page >= WORDS_PER_PAGE:
            # finish current page
            page_number = len(pages) + 1
            page_html = (
                f'<div class="word-preview-page preview-page" data-page="{page_number}">'
                '<div class="word-preview-sheet">'
                + ''.join(current_parts)
                + '</div></div>'
            )
            pages.append(page_html)
            # reset
            current_parts = []
            words_in_current_page = 0

    for obj in word_objects:
        if obj["type"] == "table_start":
            flush_paragraph()
            current_parts.append('<div class="word-preview-table-wrap">')
            current_parts.append('<table class="word-preview-table"><tbody><tr><td class="word-preview-cell">')
            in_table = True

        elif obj["type"] == "table_end":
            flush_paragraph()
            flush_table_end()
            in_table = False

        elif obj["type"] == "word":
            token_classes = ["word-preview-token"]
            if obj.get("is_heading"):
                token_classes.append("is-heading")
                # mark paragraph as a heading when the first word of the paragraph is a heading
                if not paragraph_tokens and not in_table:
                    paragraph_is_heading = True
            if obj.get("is_bold"):
                token_classes.append("is-bold")
            if obj.get("is_italic"):
                token_classes.append("is-italic")

            escaped_text = html.escape(obj["text"])
            class_name = " ".join(token_classes)
            token_html = f'<span class="{class_name}">{escaped_text}</span>'
            if text_idx in diff_indices:
                token_html = f'<span class="word-highlight">{token_html}</span>'

            if in_table:
                current_parts.append(token_html)
                current_parts.append(' ')
            else:
                paragraph_tokens.append(token_html)

            text_idx += 1
            words_in_current_page += 1
            start_new_page_if_needed()

        elif obj["type"] == "image":
            image_html = (
                '<div class="word-preview-image-block">'
                f'<img class="word-preview-inline-image" src="data:{obj.get("content_type", "image/png")};base64,{obj["src"]}" alt="Embedded document image" />'
                '</div>'
            )
            if in_table:
                current_parts.append(image_html)
            else:
                flush_paragraph()
                current_parts.append(image_html)

        elif obj["type"] == "row_end":
            if in_table:
                current_parts.append('</td></tr><tr><td class="word-preview-cell">')

        elif obj["type"] == "separator":
            if in_table:
                current_parts.append('</td><td class="word-preview-cell">')

        elif obj["type"] == "newline":
            if not in_table:
                flush_paragraph()

    # finalize
    flush_paragraph()
    if current_parts:
        page_number = len(pages) + 1
        page_html = (
            f'<div class="word-preview-page preview-page" data-page="{page_number}">'
            '<div class="word-preview-sheet">'
            + ''.join(current_parts)
            + '</div></div>'
        )
        pages.append(page_html)

    # If no pages produced (empty doc), return an empty sheet
    if not pages:
        return '<div class="word-preview-page preview-page" data-page="1"><div class="word-preview-sheet"></div></div>', 1

    return ''.join(pages), len(pages)


def _render_pdf_preview_base64(
    doc: fitz.Document,
    word_data: list[dict[str, Any]],
    diff_indices: set[int],
    max_pages: int,
    include_images: bool,
) -> list[dict[str, Any]]:
    previews: list[dict[str, Any]] = []
    total_pages = len(doc) if max_pages <= 0 else min(max_pages, len(doc))

    page_to_words: dict[int, list[dict[str, Any]]] = {}
    for idx in diff_indices:
        if idx >= len(word_data):
            continue
        word_info = word_data[idx]
        page_to_words.setdefault(word_info["page"], []).append(word_info)

    for page_num in range(total_pages):
        page = doc[page_num]
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        draw = ImageDraw.Draw(image, "RGBA")

        for word_info in page_to_words.get(page_num, []):

            x0, y0, x1, y1 = word_info["bbox"]
            x0, y0, x1, y1 = x0 * 2, y0 * 2, x1 * 2, y1 * 2
            fill_color = (255, 237, 224, 110) if word_info.get("in_table", False) else (255, 255, 0, 110)
            draw.rectangle([x0, y0, x1, y1], fill=fill_color, outline=(255, 200, 0, 200), width=2)

        page_payload: dict[str, Any] = {"page": page_num + 1}
        if include_images:
            output = BytesIO()
            image.save(output, format="JPEG", quality=55, optimize=True)
            encoded = base64.b64encode(output.getvalue()).decode("utf-8")
            page_payload["image_base64"] = encoded
            page_payload["mime_type"] = "image/jpeg"
        # include original image dimensions so the frontend can size and align pages
        page_payload["width_px"] = pix.width
        page_payload["height_px"] = pix.height

        previews.append(page_payload)

    return previews


def _safe_name(name: str) -> str:
    return Path(name).name


def _extension(filename: str) -> str:
    return Path(filename).suffix.lower().lstrip(".")


def compare_documents(doc1_name: str, doc1_bytes: bytes, doc2_name: str, doc2_bytes: bytes) -> dict[str, Any]:
    ext1 = _extension(doc1_name)
    ext2 = _extension(doc2_name)

    if ext1 not in ALLOWED_EXTENSIONS or ext2 not in ALLOWED_EXTENSIONS:
        raise ValueError("Only .pdf and .docx files are supported")

    is_pdf1 = ext1 == "pdf"
    is_pdf2 = ext2 == "pdf"

    pdf_doc1 = None
    pdf_doc2 = None

    if is_pdf1:
        text1, word_objs1, highlight_data1, pdf_doc1 = _extract_words_from_pdf(doc1_bytes)
        docx_doc1 = None
    else:
        text1, word_objs1, docx_doc1 = _extract_words_from_word(doc1_bytes)
        highlight_data1 = None

    if is_pdf2:
        text2, word_objs2, highlight_data2, pdf_doc2 = _extract_words_from_pdf(doc2_bytes)
        docx_doc2 = None
    else:
        text2, word_objs2, docx_doc2 = _extract_words_from_word(doc2_bytes)
        highlight_data2 = None

    words1 = _collect_word_tokens(word_objs1)
    words2 = _collect_word_tokens(word_objs2)

    if not words1 or not words2:
        if pdf_doc1:
            pdf_doc1.close()
        if pdf_doc2:
            pdf_doc2.close()
        raise ValueError("Could not extract text from one or both documents")

    diff1, diff2, info = _run_diff(words1, words2)

    try:
        if is_pdf1:
            highlighted_doc1 = _highlight_pdf_words(pdf_doc1, highlight_data1, diff1)
        else:
            highlighted_doc1 = _highlight_word_doc(docx_doc1, word_objs1, diff1)

        if is_pdf2:
            highlighted_doc2 = _highlight_pdf_words(pdf_doc2, highlight_data2, diff2)
        else:
            highlighted_doc2 = _highlight_word_doc(docx_doc2, word_objs2, diff2)
    finally:
        if pdf_doc1:
            pdf_doc1.close()
        if pdf_doc2:
            pdf_doc2.close()

    original_name1 = Path(_safe_name(doc1_name)).stem
    original_name2 = Path(_safe_name(doc2_name)).stem

    return {
        "doc1_output_name": f"highlighted_{original_name1}.{ext1}",
        "doc2_output_name": f"highlighted_{original_name2}.{ext2}",
        "doc1_bytes": highlighted_doc1.getvalue(),
        "doc2_bytes": highlighted_doc2.getvalue(),
        "doc1_ext": ext1,
        "doc2_ext": ext2,
        "summary": {
            **info,
            "highlighted_changes": len(diff1) + len(diff2),
            "match_rate": round((info["total_matching"] / max(info["total_words1"], info["total_words2"])) * 100, 2)
            if max(info["total_words1"], info["total_words2"]) > 0
            else 0,
        },
    }


def compare_documents_with_preview(
    doc1_name: str,
    doc1_bytes: bytes,
    doc2_name: str,
    doc2_bytes: bytes,
    max_pages: int = 0,
    include_images: bool = False,
) -> dict[str, Any]:
    ext1 = _extension(doc1_name)
    ext2 = _extension(doc2_name)

    if ext1 not in ALLOWED_EXTENSIONS or ext2 not in ALLOWED_EXTENSIONS:
        raise ValueError("Only .pdf and .docx files are supported")

    is_pdf1 = ext1 == "pdf"
    is_pdf2 = ext2 == "pdf"

    pdf_doc1 = None
    pdf_doc2 = None

    if is_pdf1:
        text1, word_objs1, highlight_data1, pdf_doc1 = _extract_words_from_pdf(doc1_bytes)
        docx_doc1 = None
    else:
        text1, word_objs1, docx_doc1 = _extract_words_from_word(doc1_bytes)
        highlight_data1 = None

    if is_pdf2:
        text2, word_objs2, highlight_data2, pdf_doc2 = _extract_words_from_pdf(doc2_bytes)
        docx_doc2 = None
    else:
        text2, word_objs2, docx_doc2 = _extract_words_from_word(doc2_bytes)
        highlight_data2 = None

    words1 = _collect_word_tokens(word_objs1)
    words2 = _collect_word_tokens(word_objs2)

    if not words1 or not words2:
        if pdf_doc1:
            pdf_doc1.close()
        if pdf_doc2:
            pdf_doc2.close()
        raise ValueError("Could not extract text from one or both documents")

    diff1, diff2, info = _run_diff(words1, words2)

    # Convert DOCX to PDF for preview rendering (for visual consistency)
    preview_pdf_doc1 = None
    preview_pdf_doc2 = None
    preview_highlight_data1 = highlight_data1
    preview_highlight_data2 = highlight_data2
    preview_diff1 = diff1
    preview_diff2 = diff2

    try:
        if not is_pdf1:
            try:
                highlighted_docx1 = _highlight_word_doc(docx_doc1, word_objs1, diff1)
                # Convert DOCX to PDF for preview
                pdf_bytes1 = _convert_docx_to_pdf(highlighted_docx1.getvalue())
                preview_pdf_doc1 = fitz.open(stream=pdf_bytes1, filetype="pdf")
                preview_highlight_data1 = []
                preview_diff1 = set()
            except Exception as ex:
                raise ValueError(f"Failed to convert Document 1 (DOCX) to PDF for preview: {str(ex)}")
        else:
            preview_pdf_doc1 = pdf_doc1

        if not is_pdf2:
            try:
                highlighted_docx2 = _highlight_word_doc(docx_doc2, word_objs2, diff2)
                # Convert DOCX to PDF for preview
                pdf_bytes2 = _convert_docx_to_pdf(highlighted_docx2.getvalue())
                preview_pdf_doc2 = fitz.open(stream=pdf_bytes2, filetype="pdf")
                preview_highlight_data2 = []
                preview_diff2 = set()
            except Exception as ex:
                raise ValueError(f"Failed to convert Document 2 (DOCX) to PDF for preview: {str(ex)}")
        else:
            preview_pdf_doc2 = pdf_doc2

        # Render both as PDF page images
        doc1_page_limit = len(preview_pdf_doc1) if max_pages <= 0 else min(max_pages, len(preview_pdf_doc1))
        preview1 = {
            "type": "pdf_images",
            "pages": _render_pdf_preview_base64(
                preview_pdf_doc1,
                preview_highlight_data1,
                preview_diff1,
                max_pages=max_pages,
                include_images=include_images,
            ),
            "images_included": include_images,
            "page_count": len(preview_pdf_doc1),
            "truncated": len(preview_pdf_doc1) > doc1_page_limit,
            "total_pages": len(preview_pdf_doc1),
        }

        doc2_page_limit = len(preview_pdf_doc2) if max_pages <= 0 else min(max_pages, len(preview_pdf_doc2))
        preview2 = {
            "type": "pdf_images",
            "pages": _render_pdf_preview_base64(
                preview_pdf_doc2,
                preview_highlight_data2,
                preview_diff2,
                max_pages=max_pages,
                include_images=include_images,
            ),
            "images_included": include_images,
            "page_count": len(preview_pdf_doc2),
            "truncated": len(preview_pdf_doc2) > doc2_page_limit,
            "total_pages": len(preview_pdf_doc2),
        }
    finally:
        # Close original documents safely
        try:
            if pdf_doc1:
                pdf_doc1.close()
        except Exception:
            pass
        try:
            if pdf_doc2:
                pdf_doc2.close()
        except Exception:
            pass
        # Close preview documents (if they're not the original PDFs) safely
        try:
            if preview_pdf_doc1 and not is_pdf1:
                preview_pdf_doc1.close()
        except Exception:
            pass
        try:
            if preview_pdf_doc2 and not is_pdf2:
                preview_pdf_doc2.close()
        except Exception:
            pass

    return {
        "summary": {
            **info,
            "highlighted_changes": len(diff1) + len(diff2),
            "match_rate": round((info["total_matching"] / max(info["total_words1"], info["total_words2"])) * 100, 2)
            if max(info["total_words1"], info["total_words2"]) > 0
            else 0,
        },
        "doc1": {"name": Path(_safe_name(doc1_name)).name, "extension": ext1, "preview": preview1},
        "doc2": {"name": Path(_safe_name(doc2_name)).name, "extension": ext2, "preview": preview2},
    }
