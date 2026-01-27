import streamlit as st
import difflib
from io import BytesIO
import fitz  # PyMuPDF
from docx import Document
import re
import os
import tempfile

st.set_page_config(page_title="Document Diff Checker", layout="wide")

st.title("📄 Document Diff Checker")
st.markdown("Upload two documents (PDF or Word) to compare and highlight their differences")

# Radio button for comparison mode
comparison_mode = st.radio(
    "Select comparison mode:",
    ["Same Format (PDF vs PDF or Word vs Word)", "Mixed Format (PDF vs Word)"],
    horizontal=True
)

# Create two columns for file uploads
col1, col2 = st.columns(2)

with col1:
    st.subheader("Document 1")
    if comparison_mode == "Same Format (PDF vs PDF or Word vs Word)":
        doc1_file = st.file_uploader("Upload first document", type=['pdf', 'docx'], key="doc1")
    else:
        doc1_file = st.file_uploader("Upload first document (PDF or Word)", type=['pdf', 'docx'], key="doc1_mixed")
    
with col2:
    st.subheader("Document 2")
    if comparison_mode == "Same Format (PDF vs PDF or Word vs Word)":
        doc2_file = st.file_uploader("Upload second document (same format as Doc 1)", type=['pdf', 'docx'], key="doc2")
    else:
        doc2_file = st.file_uploader("Upload second document (PDF or Word)", type=['pdf', 'docx'], key="doc2_mixed")

def extract_text_from_word(docx_file):
    """Extract text from Word document maintaining structure"""
    try:
        docx_file.seek(0)
        doc = Document(docx_file)
        
        text_segments = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                text_segments.append(para_text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_segments.append(' | '.join(row_texts))
        
        # Join segments with double newline to preserve structure
        extracted_text = '\n\n'.join(text_segments)
        
        return extracted_text
    except Exception as e:
        st.error(f"Error reading Word document: {str(e)}")
        return None

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF with enhanced table detection"""
    try:
        pdf_file.seek(0)
        doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
        
        all_words = []
        full_text_parts = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Try to detect tables
            tables = page.find_tables()
            table_rects = []
            
            if tables:
                for table in tables:
                    table_rects.append(table.bbox)
            
            # Get word-level data with coordinates
            words_data = page.get_text("words")
            
            for word_info in words_data:
                x0, y0, x1, y1, word_text = word_info[:5]
                
                if word_text.strip():
                    # Check if word is in a table
                    in_table = False
                    for table_rect in table_rects:
                        if (x0 >= table_rect[0] and x1 <= table_rect[2] and 
                            y0 >= table_rect[1] and y1 <= table_rect[3]):
                            in_table = True
                            break
                    
                    all_words.append({
                        'text': word_text.strip(),
                        'bbox': [x0, y0, x1, y1],
                        'page': page_num,
                        'in_table': in_table
                    })
                    full_text_parts.append(word_text.strip())
        
        # Join with spaces
        full_text = ' '.join(full_text_parts)
        
        return full_text, all_words, doc
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return None, None, None

# ============================================================================
# SAME FORMAT COMPARISON FUNCTIONS (FROM FIRST CODE FILE)
# ============================================================================

def normalize_word_same_format(word):
    """Normalize word by removing ALL punctuation and converting to lowercase for comparison"""
    # Remove ALL punctuation including brackets, quotes, dashes, periods, etc.
    # This ensures words are compared purely on alphanumeric content
    import string
    # Create a translation table that removes all punctuation
    translator = str.maketrans('', '', string.punctuation)
    # Remove all punctuation and convert to lowercase
    normalized = word.translate(translator).lower()
    # Also handle special unicode quotes and dashes
    normalized = normalized.replace('"', '').replace('"', '').replace(''', '').replace(''', '')
    normalized = normalized.replace('–', '').replace('—', '')
    return normalized

def normalize_for_comparison_same_format(text):
    """Normalize text for comparison - removes ALL punctuation differences"""
    # Split into words
    words = text.split()
    # Normalize each word and filter out empty strings (pure punctuation)
    normalized = []
    for w in words:
        norm_word = normalize_word_same_format(w)
        if norm_word:  # Only include if there's actual content left after removing punctuation
            normalized.append(norm_word)
    return normalized

def find_word_differences_optimized_same_format(text1, text2):
    """
    Optimized word-level difference detection with better alignment handling.
    Uses content-based matching to avoid drift in long documents.
    Excludes punctuation-only differences.
    (FROM FIRST CODE FILE)
    """
    # Split into words
    words1 = text1.split()
    words2 = text2.split()
    
    # Create normalized versions for comparison (excluding punctuation)
    normalized1 = normalize_for_comparison_same_format(text1)
    normalized2 = normalize_for_comparison_same_format(text2)
    
    # Use SequenceMatcher with optimized settings
    matcher = difflib.SequenceMatcher(None, normalized1, normalized2, autojunk=False)
    
    # Get the opcodes
    opcodes = matcher.get_opcodes()
    
    # Sets to store indices of different words in NORMALIZED space
    diff_indices_norm1 = set()
    diff_indices_norm2 = set()
    
    # Create mapping from normalized to original indices
    norm_to_orig1 = []
    norm_to_orig2 = []
    
    for i, word in enumerate(words1):
        if normalize_word_same_format(word):
            norm_to_orig1.append(i)
    
    for i, word in enumerate(words2):
        if normalize_word_same_format(word):
            norm_to_orig2.append(i)
    
    # Process each operation
    for tag, i1, i2, j1, j2 in opcodes:
        if tag == 'equal':
            # Words match - don't highlight
            continue
        elif tag == 'replace':
            # Check if this is a true content difference
            range1_words = normalized1[i1:i2]
            range2_words = normalized2[j1:j2]
            
            # Only mark as different if normalized content actually differs
            if range1_words != range2_words:
                diff_indices_norm1.update(range(i1, i2))
                diff_indices_norm2.update(range(j1, j2))
        elif tag == 'delete':
            # Words only in doc1
            diff_indices_norm1.update(range(i1, i2))
        elif tag == 'insert':
            # Words only in doc2
            diff_indices_norm2.update(range(j1, j2))
    
    # Map back to original word indices
    diff_indices1 = set()
    diff_indices2 = set()
    
    for norm_idx in diff_indices_norm1:
        if norm_idx < len(norm_to_orig1):
            diff_indices1.add(norm_to_orig1[norm_idx])
    
    for norm_idx in diff_indices_norm2:
        if norm_idx < len(norm_to_orig2):
            diff_indices2.add(norm_to_orig2[norm_idx])
    
    # CRITICAL: Remove false positives where both docs have same word at same position
    # This happens when alignment drifts
    validated_diff1 = set()
    validated_diff2 = set()
    
    # Create a reverse mapping to find normalized positions from original positions
    orig_to_norm1 = {}
    for norm_idx, orig_idx in enumerate(norm_to_orig1):
        orig_to_norm1[orig_idx] = norm_idx
    
    orig_to_norm2 = {}
    for norm_idx, orig_idx in enumerate(norm_to_orig2):
        orig_to_norm2[orig_idx] = norm_idx
    
    # Validate diff_indices1
    for orig_idx in diff_indices1:
        if orig_idx < len(words1):
            word1_norm = normalize_word_same_format(words1[orig_idx])
            
            # Find corresponding position in doc2
            if orig_idx in orig_to_norm1:
                norm_idx1 = orig_to_norm1[orig_idx]
                
                # Check if this normalized index has a corresponding position in doc2
                # by looking at the opcodes to find the mapping
                is_truly_different = True
                
                for tag, i1, i2, j1, j2 in opcodes:
                    if tag == 'equal' and i1 <= norm_idx1 < i2:
                        # This word is in an equal block, find its pair
                        offset = norm_idx1 - i1
                        norm_idx2 = j1 + offset
                        
                        if norm_idx2 < len(norm_to_orig2):
                            orig_idx2 = norm_to_orig2[norm_idx2]
                            if orig_idx2 < len(words2):
                                word2_norm = normalize_word_same_format(words2[orig_idx2])
                                # If words are actually the same, don't highlight
                                if word1_norm == word2_norm:
                                    is_truly_different = False
                                    break
                
                if is_truly_different:
                    validated_diff1.add(orig_idx)
            else:
                validated_diff1.add(orig_idx)
    
    # Validate diff_indices2
    for orig_idx in diff_indices2:
        if orig_idx < len(words2):
            word2_norm = normalize_word_same_format(words2[orig_idx])
            
            # Find corresponding position in doc1
            if orig_idx in orig_to_norm2:
                norm_idx2 = orig_to_norm2[orig_idx]
                
                is_truly_different = True
                
                for tag, i1, i2, j1, j2 in opcodes:
                    if tag == 'equal' and j1 <= norm_idx2 < j2:
                        # This word is in an equal block, find its pair
                        offset = norm_idx2 - j1
                        norm_idx1 = i1 + offset
                        
                        if norm_idx1 < len(norm_to_orig1):
                            orig_idx1 = norm_to_orig1[norm_idx1]
                            if orig_idx1 < len(words1):
                                word1_norm = normalize_word_same_format(words1[orig_idx1])
                                # If words are actually the same, don't highlight
                                if word1_norm == word2_norm:
                                    is_truly_different = False
                                    break
                
                if is_truly_different:
                    validated_diff2.add(orig_idx)
            else:
                validated_diff2.add(orig_idx)
    
    diff_indices1 = validated_diff1
    diff_indices2 = validated_diff2
    
    # Calculate statistics based on normalized words
    total_matching_words = sum(i2 - i1 for tag, i1, i2, _, _ in opcodes if tag == 'equal')
    total_equal_blocks = sum(1 for tag, _, _, _, _ in opcodes if tag == 'equal')
    
    sync_info = {
        'total_matching': total_matching_words,
        'total_words1': len(words1),
        'total_words2': len(words2),
        'equal_blocks': total_equal_blocks,
        'diff_words1': len(diff_indices1),
        'diff_words2': len(diff_indices2)
    }
    
    return diff_indices1, diff_indices2, words1, words2, sync_info

# ============================================================================
# MIXED FORMAT COMPARISON FUNCTIONS (FROM SECOND CODE FILE)
# ============================================================================

def normalize_word_mixed_format(word):
    """
    Enhanced word normalization for mixed format comparison.
    More conservative normalization to handle PDF vs Word differences.
    (FROM SECOND CODE FILE)
    """
    import string
    import unicodedata
    
    # Convert to lowercase
    word = word.lower()
    
    # Normalize unicode characters
    word = unicodedata.normalize('NFKD', word)
    
    # Remove only the most problematic punctuation, preserve structure
    # Remove quotes and apostrophes
    word = word.replace('"', '').replace("'", '').replace('''', '').replace(''', '')
    word = word.replace('"', '').replace('"', '')
    
    # Handle dashes - replace with space to preserve word separation
    word = word.replace('–', ' ').replace('—', ' ')
    
    # Remove other punctuation but preserve letters and numbers
    # Be more selective about what we remove
    word = re.sub(r'[^\w\s-]', '', word)  # Keep letters, numbers, spaces, and hyphens
    
    # Handle common OCR errors, but be conservative
    if len(word) == 1:  # Only for single characters
        word = word.replace('0', 'o')  # Zero to letter o
        word = word.replace('1', 'l')  # One to letter l
        word = word.replace('5', 's')  # Five to letter s
    
    # Remove extra whitespace
    word = word.strip()
    
    return word

def normalize_mixed_format_text(text):
    """
    Enhanced normalization specifically for PDF vs Word comparison.
    Handles the structural differences between PDF and Word formats with conservative approach.
    (FROM SECOND CODE FILE)
    """
    # Convert to lowercase
    text = text.lower()
    
    # Handle PDF-specific issues - but be very conservative
    # Only normalize excessive whitespace, preserve structure
    text = re.sub(r'\s{3,}', '  ', text)  # Replace 3+ spaces with 2 spaces
    text = re.sub(r'\n{3,}', '\n\n', text)  # Replace 3+ newlines with 2 newlines
    
    # Very conservative hyphenation handling - only handle clear cases
    text = re.sub(r'(\w+)-\s+([a-z]+)', r'\1\2', text)  # Join hyphenated words with space after
    
    # Normalize common spacing issues around punctuation
    text = re.sub(r'\s+([,.;:!?])', r'\1', text)  # Remove spaces before punctuation
    text = re.sub(r'([,.;:!?])\s+', r'\1 ', text)  # Ensure single space after punctuation
    
    # Remove excessive whitespace at start/end
    text = text.strip()
    
    return text

def find_differences_mixed_format(text1, text2, is_pdf1_word2):
    """
    Enhanced difference detection specifically for PDF vs Word comparison.
    Handles format-specific text extraction differences with conservative approach.
    (FROM SECOND CODE FILE)
    """
    # Apply mixed format normalization
    norm_text1 = normalize_mixed_format_text(text1)
    norm_text2 = normalize_mixed_format_text(text2)
    
    # Further normalize words for comparison
    words1 = norm_text1.split()
    words2 = norm_text2.split()
    
    # Create even more normalized versions for comparison
    norm_words1 = [normalize_word_mixed_format(w) for w in words1 if normalize_word_mixed_format(w)]
    norm_words2 = [normalize_word_mixed_format(w) for w in words2 if normalize_word_mixed_format(w)]
    
    # Use SequenceMatcher with very conservative settings for mixed format
    matcher = difflib.SequenceMatcher(None, norm_words1, norm_words2, autojunk=True)
    
    # Get opcodes
    opcodes = matcher.get_opcodes()
    
    # Find differences with very conservative validation
    diff_indices1 = set()
    diff_indices2 = set()
    
    for tag, i1, i2, j1, j2 in opcodes:
        if tag == 'equal':
            continue
        elif tag == 'replace':
            # Very conservative validation for replacement
            range1 = ' '.join(norm_words1[i1:i2])
            range2 = ' '.join(norm_words2[j1:j2])
            
            # Additional check: only mark as different if content is significantly different
            similarity = difflib.SequenceMatcher(None, range1, range2).ratio()
            if similarity < 0.7:  # Only mark as different if less than 70% similar (more conservative)
                # Additional validation: check if the words are truly different
                words1_in_range = norm_words1[i1:i2]
                words2_in_range = norm_words2[j1:j2]
                
                # Only mark as different if there's a significant content difference
                # Check if the core meaning is different
                if len(words1_in_range) != len(words2_in_range) or any(w1 != w2 for w1, w2 in zip(words1_in_range, words2_in_range)):
                    # Map back to original indices with very conservative mapping
                    orig_start1 = sum(len(w.split()) for w in norm_words1[:i1])
                    orig_end1 = orig_start1 + sum(len(w.split()) for w in norm_words1[i1:i2])
                    
                    orig_start2 = sum(len(w.split()) for w in norm_words2[:j1])
                    orig_end2 = orig_start2 + sum(len(w.split()) for w in norm_words2[j1:j2])
                    
                    # Only add indices if there's a real content difference
                    if similarity < 0.6:  # Even more conservative for final decision
                        for i in range(orig_start1, orig_end1):
                            diff_indices1.add(i)
                        for i in range(orig_start2, orig_end2):
                            diff_indices2.add(i)
        
        elif tag == 'delete':
            # Words only in doc1 - be very conservative
            if (i2 - i1) > 0:  # Only if there's actually content missing
                orig_start = sum(len(w.split()) for w in norm_words1[:i1])
                orig_end = orig_start + sum(len(w.split()) for w in norm_words1[i1:i2])
                for i in range(orig_start, orig_end):
                    diff_indices1.add(i)
        
        elif tag == 'insert':
            # Words only in doc2 - be very conservative  
            if (j2 - j1) > 0:  # Only if there's actually content added
                orig_start = sum(len(w.split()) for w in norm_words2[:j1])
                orig_end = orig_start + sum(len(w.split()) for w in norm_words2[j1:j2])
                for i in range(orig_start, orig_end):
                    diff_indices2.add(i)
    
    # Post-processing: Remove potential false positives
    # If too many differences are found, relax the criteria
    total_words1 = len(text1.split())
    total_words2 = len(text2.split())
    
    diff_ratio1 = len(diff_indices1) / total_words1 if total_words1 > 0 else 0
    diff_ratio2 = len(diff_indices2) / total_words2 if total_words2 > 0 else 0
    
    # If more than 30% of words are marked as different, it's likely false positives
    # This suggests the documents are mostly the same but formatting differences
    if diff_ratio1 > 0.3 or diff_ratio2 > 0.3:
        # Clear the differences and use a much more lenient approach
        diff_indices1.clear()
        diff_indices2.clear()
        
        # Re-run with very lenient settings - only mark obvious differences
        for tag, i1, i2, j1, j2 in opcodes:
            if tag == 'equal':
                continue
            elif tag == 'replace':
                range1 = ' '.join(norm_words1[i1:i2])
                range2 = ' '.join(norm_words2[j1:j2])
                similarity = difflib.SequenceMatcher(None, range1, range2).ratio()
                
                # Only mark as different if very obviously different (less than 40% similar)
                if similarity < 0.4:
                    orig_start1 = sum(len(w.split()) for w in norm_words1[:i1])
                    orig_end1 = orig_start1 + sum(len(w.split()) for w in norm_words1[i1:i2])
                    
                    orig_start2 = sum(len(w.split()) for w in norm_words2[:j1])
                    orig_end2 = orig_start2 + sum(len(w.split()) for w in norm_words2[j1:j2])
                    
                    for i in range(orig_start1, orig_end1):
                        diff_indices1.add(i)
                    for i in range(orig_start2, orig_end2):
                        diff_indices2.add(i)
    
    # Calculate statistics
    total_matching = sum(i2 - i1 for tag, i1, i2, _, _ in opcodes if tag == 'equal')
    
    sync_info = {
        'total_matching': total_matching,
        'total_words1': total_words1,
        'total_words2': total_words2,
        'diff_words1': len(diff_indices1),
        'diff_words2': len(diff_indices2)
    }
    
    return diff_indices1, diff_indices2, text1.split(), text2.split(), sync_info

# ============================================================================
# COMMON FUNCTIONS (USED BY BOTH)
# ============================================================================

def create_html_diff(text1, text2, diff_indices1, diff_indices2):
    """Create HTML with highlighted differences"""
    def highlight_text(text, diff_indices):
        words = text.split()
        html_parts = []
        
        for i, word in enumerate(words):
            if i in diff_indices:
                html_parts.append(f'<span class="highlight">{word}</span>')
            else:
                html_parts.append(word)
        
        return ' '.join(html_parts)
    
    html1 = highlight_text(text1, diff_indices1)
    html2 = highlight_text(text2, diff_indices2)
    
    return html1, html2

def highlight_pdf_words(doc, word_data, diff_indices):
    """Highlight specific words in PDF based on indices with enhanced table support"""
    highlighted_doc = fitz.open()
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        new_page = highlighted_doc.new_page(width=page.rect.width, height=page.rect.height)
        new_page.show_pdf_page(new_page.rect, doc, page_num)
        
        # Highlight words on this page
        for word_idx in diff_indices:
            if word_idx < len(word_data):
                word_info = word_data[word_idx]
                if word_info['page'] == page_num:
                    bbox = word_info['bbox']
                    rect = fitz.Rect(bbox[0], bbox[1], bbox[2], bbox[3])
                    
                    try:
                        # Use different color for table content
                        if word_info.get('in_table', False):
                            highlight = new_page.add_highlight_annot(rect)
                            # Light orange/peach color (RGB: 255, 200, 150)
                            highlight.set_colors(stroke=[1.0, 0.93, 0.88])
                        else:
                            highlight = new_page.add_highlight_annot(rect)
                            highlight.set_colors(stroke=fitz.utils.getColor("yellow"))
                        highlight.update()
                    except:
                        pass
    
    return highlighted_doc

def highlight_word_doc(docx_file, extracted_text, diff_indices):
    """
    Highlight words in Word document with precise run-level matching.
    """
    from docx.enum.text import WD_COLOR_INDEX
    
    docx_file.seek(0)
    doc = Document(docx_file)
    
    # Get all words from extracted text
    extracted_words = extracted_text.replace('\n\n', ' ').split()
    
    # Build a map of word positions to track what we've processed
    word_idx = 0
    
    # Process regular paragraphs
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if not para_text:
            continue
        
        para_words = para_text.split()
        
        # Match this paragraph's words to extracted words
        para_start_idx = word_idx
        
        # Check if current paragraph words match extracted words at current position
        matches = True
        for i, pword in enumerate(para_words):
            if word_idx + i >= len(extracted_words):
                matches = False
                break
            if normalize_word_same_format(pword) != normalize_word_same_format(extracted_words[word_idx + i]):
                matches = False
                break
        
        if not matches:
            # Try to find where this paragraph appears in extracted words
            # This handles cases where paragraphs might be out of sync
            for search_offset in range(max(0, word_idx - 10), min(len(extracted_words), word_idx + 50)):
                temp_matches = True
                for i, pword in enumerate(para_words):
                    if search_offset + i >= len(extracted_words):
                        temp_matches = False
                        break
                    if normalize_word_same_format(pword) != normalize_word_same_format(extracted_words[search_offset + i]):
                        temp_matches = False
                        break
                if temp_matches:
                    word_idx = search_offset
                    para_start_idx = search_offset
                    break
        
        # Now highlight runs in this paragraph
        run_word_position = para_start_idx
        
        for run in para.runs:
            run_text = run.text
            if not run_text:
                continue
            
            run_words = run_text.split()
            if not run_words:
                continue
            
            # Check if ANY word in this run should be highlighted
            should_highlight = False
            for i in range(len(run_words)):
                if (run_word_position + i) in diff_indices:
                    should_highlight = True
                    break
            
            if should_highlight:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            run_word_position += len(run_words)
        
        # Move global word index forward
        word_idx += len(para_words)
    
    # Process tables with improved logic
    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            
            # First pass: collect all cell texts for this row
            for cell in row.cells:
                cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if cell_text:
                    row_cells.append(cell_text)
            
            # Check if we can find this row in extracted text
            if not row_cells:
                continue
            
            # The row might be represented as "cell1 | cell2 | cell3" in extracted text
            expected_row_text = ' | '.join(row_cells)
            
            # Try to find this in extracted text starting from current position
            row_found = False
            search_range = ' '.join(extracted_words[word_idx:min(word_idx + 200, len(extracted_words))])
            
            if expected_row_text in search_range:
                row_found = True
            
            # Process each cell
            for cell in row.cells:
                for para in cell.paragraphs:
                    para_text = para.text.strip()
                    if not para_text:
                        continue
                    
                    para_words = para_text.split()
                    cell_start_idx = word_idx
                    
                    # Verify alignment
                    matches = True
                    for i, cword in enumerate(para_words):
                        if word_idx + i >= len(extracted_words):
                            matches = False
                            break
                        if normalize_word_same_format(cword) != normalize_word_same_format(extracted_words[word_idx + i]):
                            matches = False
                            break
                    
                    # Highlight runs in this cell paragraph
                    run_word_position = cell_start_idx if matches else word_idx
                    
                    for run in para.runs:
                        run_text = run.text
                        if not run_text:
                            continue
                        
                        run_words = run_text.split()
                        if not run_words:
                            continue
                        
                        # Check if ANY word in this run should be highlighted
                        should_highlight = False
                        for i in range(len(run_words)):
                            if (run_word_position + i) in diff_indices:
                                should_highlight = True
                                break
                        
                        if should_highlight:
                            # Use bright green for table differences
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                        
                        run_word_position += len(run_words)
                    
                    if matches:
                        word_idx += len(para_words)
                
                # Account for cell separator in extracted text (the " | " part)
                if word_idx < len(extracted_words) and cell != row.cells[-1]:
                    # Skip the separator marker if present
                    if word_idx + 1 < len(extracted_words) and extracted_words[word_idx] == '|':
                        word_idx += 1
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ============================================================================
# COMPARISON ORCHESTRATION FUNCTIONS
# ============================================================================

def compare_same_format(doc1_file, doc2_file):
    """
    Compare documents of the same format (PDF vs PDF or Word vs Word)
    Uses the original logic from the first code file
    """
    # Detect file types
    is_pdf1 = doc1_file.name.endswith('.pdf')
    is_pdf2 = doc2_file.name.endswith('.pdf')
    
    with st.spinner("Extracting text from documents..."):
        if is_pdf1:
            text1, word_data1, pdf_doc1 = extract_text_from_pdf(doc1_file)
        else:
            text1 = extract_text_from_word(doc1_file)
            word_data1 = None
            pdf_doc1 = None
        
        if is_pdf2:
            text2, word_data2, pdf_doc2 = extract_text_from_pdf(doc2_file)
        else:
            text2 = extract_text_from_word(doc2_file)
            word_data2 = None
            pdf_doc2 = None
    
    if not text1 or not text2:
        return None
    
    with st.spinner("Finding differences (excluding punctuation)..."):
        diff_indices1, diff_indices2, words1, words2, sync_info = find_word_differences_optimized_same_format(text1, text2)
        html1, html2 = create_html_diff(text1, text2, diff_indices1, diff_indices2)
    
    with st.spinner("Generating highlighted documents..."):
        # Generate highlighted versions
        if is_pdf1:
            highlighted_doc1 = highlight_pdf_words(pdf_doc1, word_data1, diff_indices1)
            pdf1_bytes = BytesIO()
            highlighted_doc1.save(pdf1_bytes)
            pdf1_bytes.seek(0)
            highlighted_doc1.close()
            pdf_doc1.close()
        else:
            pdf1_bytes = highlight_word_doc(doc1_file, text1, diff_indices1)
        
        if is_pdf2:
            highlighted_doc2 = highlight_pdf_words(pdf_doc2, word_data2, diff_indices2)
            pdf2_bytes = BytesIO()
            highlighted_doc2.save(pdf2_bytes)
            pdf2_bytes.seek(0)
            highlighted_doc2.close()
            pdf_doc2.close()
        else:
            pdf2_bytes = highlight_word_doc(doc2_file, text2, diff_indices2)
    
    return {
        'text1': text1,
        'text2': text2,
        'diff_indices1': diff_indices1,
        'diff_indices2': diff_indices2,
        'words1': words1,
        'words2': words2,
        'html1': html1,
        'html2': html2,
        'pdf1_bytes': pdf1_bytes,
        'pdf2_bytes': pdf2_bytes,
        'is_pdf1': is_pdf1,
        'is_pdf2': is_pdf2,
        'sync_info': sync_info,
        'is_mixed_format': False
    }

def compare_mixed_format(doc1_file, doc2_file):
    """
    Compare PDF vs Word document with enhanced mixed format handling.
    Uses the logic from the second code file
    """
    # Detect which is PDF and which is Word
    is_pdf1 = doc1_file.name.endswith('.pdf')
    is_pdf2 = doc2_file.name.endswith('.pdf')
    
    # Determine PDF and Word files
    if is_pdf1 and not is_pdf2:
        pdf_file = doc1_file
        word_file = doc2_file
        pdf_is_doc1 = True
    elif not is_pdf1 and is_pdf2:
        pdf_file = doc2_file
        word_file = doc1_file
        pdf_is_doc1 = False
    else:
        st.error("For mixed format comparison, please upload one PDF and one Word document.")
        return None
    
    with st.spinner("Extracting text from documents..."):
        # Extract from PDF
        text_pdf, word_data_pdf, pdf_doc = extract_text_from_pdf(pdf_file)
        
        # Extract from Word
        text_word = extract_text_from_word(word_file)
    
    if not text_pdf or not text_word:
        return None
    
    # Assign texts based on which is doc1
    if pdf_is_doc1:
        text1, text2 = text_pdf, text_word
        word_data1 = word_data_pdf
        word_data2 = None
    else:
        text1, text2 = text_word, text_pdf
        word_data1 = None
        word_data2 = word_data_pdf
    
    with st.spinner("Finding differences with enhanced mixed format handling..."):
        # Use enhanced mixed format comparison
        diff_indices1, diff_indices2, words1, words2, sync_info = find_differences_mixed_format(text1, text2, pdf_is_doc1)
        html1, html2 = create_html_diff(text1, text2, diff_indices1, diff_indices2)
    
    with st.spinner("Generating highlighted documents..."):
        # Highlight PDF
        if pdf_is_doc1:
            highlighted_pdf = highlight_pdf_words(pdf_doc, word_data_pdf, diff_indices1)
            pdf_bytes = BytesIO()
            highlighted_pdf.save(pdf_bytes)
            pdf_bytes.seek(0)
            highlighted_pdf.close()
            
            # Highlight Word
            word_bytes = highlight_word_doc(word_file, text_word, diff_indices2)
            
            doc1_bytes = pdf_bytes
            doc2_bytes = word_bytes
        else:
            # Highlight Word
            word_bytes = highlight_word_doc(word_file, text_word, diff_indices1)
            
            # Highlight PDF
            highlighted_pdf = highlight_pdf_words(pdf_doc, word_data_pdf, diff_indices2)
            pdf_bytes = BytesIO()
            highlighted_pdf.save(pdf_bytes)
            pdf_bytes.seek(0)
            highlighted_pdf.close()
            
            doc1_bytes = word_bytes
            doc2_bytes = pdf_bytes
        
        pdf_doc.close()
    
    return {
        'text1': text1,
        'text2': text2,
        'diff_indices1': diff_indices1,
        'diff_indices2': diff_indices2,
        'words1': words1,
        'words2': words2,
        'html1': html1,
        'html2': html2,
        'pdf1_bytes': doc1_bytes,
        'pdf2_bytes': doc2_bytes,
        'is_pdf1': pdf_is_doc1,
        'is_pdf2': not pdf_is_doc1,
        'sync_info': sync_info,
        'is_mixed_format': True
    }

# CSS for styling
st.markdown("""
<style>
    .diff-container {
        font-family: 'Courier New', monospace;
        font-size: 13px;
        line-height: 1.8;
        padding: 20px;
        background-color: #ffffff;
        border-radius: 5px;
        max-height: 700px;
        overflow-y: auto;
        border: 1px solid #ddd;
        white-space: pre-wrap;
        word-wrap: break-word;
    }
    .highlight {
        background-color: #ffff00;
        padding: 1px 2px;
        font-weight: bold;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'comparison_done' not in st.session_state:
    st.session_state.comparison_done = False
if 'results' not in st.session_state:
    st.session_state.results = None
if 'last_mode' not in st.session_state:
    st.session_state.last_mode = None

# Process and display differences
if doc1_file and doc2_file:
    # Check if mode or files changed
    current_files = (doc1_file.name, doc2_file.name)
    if ('last_files' not in st.session_state or 
        st.session_state.last_files != current_files or
        st.session_state.last_mode != comparison_mode):
        st.session_state.comparison_done = False
        st.session_state.last_files = current_files
        st.session_state.last_mode = comparison_mode
    
    if not st.session_state.comparison_done:
        # Validate file types based on mode
        is_pdf1 = doc1_file.name.endswith('.pdf')
        is_pdf2 = doc2_file.name.endswith('.pdf')
        
        if comparison_mode == "Same Format (PDF vs PDF or Word vs Word)":
            # Check if both files are same format
            if (is_pdf1 and is_pdf2) or (not is_pdf1 and not is_pdf2):
                results = compare_same_format(doc1_file, doc2_file)
                if results:
                    st.session_state.results = results
                    st.session_state.comparison_done = True
            else:
                st.error("⚠️ For same format comparison, both documents must be the same type (both PDF or both Word).")
        
        else:  # Mixed Format mode
            # Check if files are different formats
            if (is_pdf1 and not is_pdf2) or (not is_pdf1 and is_pdf2):
                results = compare_mixed_format(doc1_file, doc2_file)
                if results:
                    st.session_state.results = results
                    st.session_state.comparison_done = True
            else:
                st.error("⚠️ For mixed format comparison, please upload one PDF and one Word document.")
    
    # Display results
    if st.session_state.results:
        results = st.session_state.results
        
        st.success("✅ Comparison complete!")
        
        # Display statistics
        sync_info = results['sync_info']
        match_percentage = (sync_info['total_matching'] / max(sync_info['total_words1'], sync_info['total_words2'])) * 100
        
        # Show comparison method used
        if results.get('is_mixed_format', False):
            st.info("🔄 Mixed format comparison algorithm used (PDF vs Word)")
        else:
            st.info("📄 Same format comparison algorithm used")
        
        st.info(f"📊 **Alignment**: {sync_info['total_matching']} matching words out of {max(sync_info['total_words1'], sync_info['total_words2'])} ({match_percentage:.1f}% match) • Punctuation differences excluded")
        
        col_stat1, col_stat2, col_stat3 = st.columns(3)
        with col_stat1:
            st.metric("Words in Doc 1", sync_info['total_words1'])
            st.metric("Different in Doc 1", sync_info['diff_words1'])
        with col_stat2:
            st.metric("Words in Doc 2", sync_info['total_words2'])
            st.metric("Different in Doc 2", sync_info['diff_words2'])
        with col_stat3:
            similarity = (sync_info['total_matching'] / max(sync_info['total_words1'], sync_info['total_words2'])) * 100
            st.metric("Match Rate", f"{similarity:.1f}%")
            st.metric("Sync Blocks", sync_info.get('equal_blocks', 0))
        
        # Download buttons
        st.markdown("### Download Highlighted Documents")
        col_dl1, col_dl2 = st.columns(2)
        
        with col_dl1:
            file_ext1 = 'pdf' if results.get('is_pdf1', False) else 'docx'
            st.download_button(
                label=f"⬇️ Download Doc 1 (Highlighted .{file_ext1})",
                data=results['pdf1_bytes'].getvalue(),
                file_name=f"doc1_highlighted.{file_ext1}",
                mime="application/pdf" if results.get('is_pdf1', False) else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        with col_dl2:
            file_ext2 = 'pdf' if results.get('is_pdf2', False) else 'docx'
            st.download_button(
                label=f"⬇️ Download Doc 2 (Highlighted .{file_ext2})",
                data=results['pdf2_bytes'].getvalue(),
                file_name=f"doc2_highlighted.{file_ext2}",
                mime="application/pdf" if results.get('is_pdf2', False) else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        # Display text comparison
        st.markdown("### Text Comparison Preview")
        st.markdown("🟡 **Yellow highlight** = Regular text differences | 🟠 **Light Orange highlight** (PDF) = Table differences | 🟢 **Green highlight** (Word) = Table differences | Punctuation differences excluded")
        
        col_diff1, col_diff2 = st.columns(2)
        
        with col_diff1:
            st.markdown("**Document 1**")
            st.markdown(f'<div class="diff-container">{results["html1"]}</div>', unsafe_allow_html=True)
        
        with col_diff2:
            st.markdown("**Document 2**")
            st.markdown(f'<div class="diff-container">{results["html2"]}</div>', unsafe_allow_html=True)
        
        # Sample differences
        with st.expander("📋 View Sample Differences (First 20)"):
            col_s1, col_s2 = st.columns(2)
            with col_s1:
                st.markdown(f"**Different words in Doc 1: {len(results['diff_indices1'])} total**")
                sample_indices1 = sorted(list(results['diff_indices1']))[:20]
                for idx in sample_indices1:
                    if idx < len(results['words1']):
                        word1 = results['words1'][idx]
                        word2 = results['words2'][idx] if idx < len(results['words2']) else "N/A"
                        st.text(f"Pos {idx}: '{word1}' vs '{word2}'")
            with col_s2:
                st.markdown(f"**Different words in Doc 2: {len(results['diff_indices2'])} total**")
                sample_indices2 = sorted(list(results['diff_indices2']))[:20]
                for idx in sample_indices2:
                    if idx < len(results['words2']):
                        word2 = results['words2'][idx]
                        word1 = results['words1'][idx] if idx < len(results['words1']) else "N/A"
                        st.text(f"Pos {idx}: '{word2}' vs '{word1}'")

else:
    st.info("👆 Please upload both documents to begin comparison")

# Footer
st.markdown("---")
st.markdown("💡 **Precise word-by-word comparison** - Highlights only the words that actually differ between documents (punctuation differences excluded)")
st.markdown("🔸 **Same Format Mode**: Uses optimized algorithm for PDF-to-PDF or Word-to-Word comparisons")
st.markdown("🔸 **Mixed Format Mode**: Uses enhanced algorithm specifically designed for PDF-to-Word comparisons")
st.markdown("🔸 For PDFs: Yellow = regular text differences, Light Orange = table differences")
st.markdown("🔸 For Word docs: Yellow = regular text differences, Green = table differences")